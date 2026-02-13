"""Word document adapter for DocumentTool."""

from __future__ import annotations

import asyncio
import io
import logging
from collections.abc import Sequence
from typing import Any

from fireflyframework_genai.tools.base import GuardProtocol

from firefly_dworkers.tools.document.base import DocumentTool
from firefly_dworkers.tools.document.models import (
    DocumentData,
    DocumentOperation,
    ParagraphData,
    SectionSpec,
)
from firefly_dworkers.tools.registry import tool_registry

try:
    import docx
    from docx.shared import Inches, Pt, RGBColor

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


def _require_docx() -> None:
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx required: pip install firefly-dworkers[document]")


@tool_registry.register("word", category="document")
class WordTool(DocumentTool):
    """Read, create, and modify Word (.docx) documents using python-docx."""

    def __init__(
        self,
        *,
        timeout: float = 60.0,
        guards: Sequence[GuardProtocol] = (),
    ) -> None:
        super().__init__(
            "word",
            description="Read, create, and modify Word (.docx) documents.",
            timeout=timeout,
            guards=guards,
        )

    async def _read_document(self, source: str) -> DocumentData:
        _require_docx()
        return await asyncio.to_thread(self._read_sync, source)

    async def _create_document(self, title: str, sections: list[SectionSpec]) -> bytes:
        _require_docx()
        return await asyncio.to_thread(self._create_sync, title, sections)

    async def _modify_document(self, source: str, operations: list[DocumentOperation]) -> bytes:
        _require_docx()
        return await asyncio.to_thread(self._modify_sync, source, operations)

    # -- Sync implementations --

    def _read_sync(self, source: str) -> DocumentData:
        doc = docx.Document(source)
        paragraphs = []
        styles: set[str] = set()

        for para in doc.paragraphs:
            is_heading = para.style.name.startswith("Heading")
            heading_level = 0
            if is_heading:
                try:
                    heading_level = int(para.style.name.split()[-1])
                except (ValueError, IndexError):
                    heading_level = 1
            paragraphs.append(
                ParagraphData(
                    text=para.text,
                    style=para.style.name,
                    is_heading=is_heading,
                    heading_level=heading_level,
                )
            )
            styles.add(para.style.name)

        title = ""
        for p in paragraphs:
            if p.text.strip():
                title = p.text.strip()
                break

        return DocumentData(
            title=title,
            paragraphs=paragraphs,
            styles=sorted(styles),
        )

    def _create_sync(self, title: str, sections: list[SectionSpec]) -> bytes:
        doc = docx.Document()

        if title:
            doc.add_heading(title, level=0)

        for section in sections:
            if section.page_break_before:
                doc.add_page_break()

            if section.heading:
                para = doc.add_heading(section.heading, level=section.heading_level)
                if section.heading_style:
                    self._apply_text_style(para, section.heading_style)

            if section.content:
                para = doc.add_paragraph(section.content)
                if section.body_style:
                    self._apply_text_style(para, section.body_style)

            for point in section.bullet_points:
                doc.add_paragraph(point, style="List Bullet")

            for item in section.numbered_list:
                doc.add_paragraph(item, style="List Number")

            if section.callout:
                self._add_callout(doc, section.callout)

            if section.table:
                headers = section.table.headers
                rows = section.table.rows
                if headers:
                    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
                    tbl.style = "Table Grid"
                    for i, header in enumerate(headers):
                        tbl.cell(0, i).text = header
                    for r, row in enumerate(rows):
                        for c, val in enumerate(row):
                            if c < len(headers):
                                tbl.cell(r + 1, c).text = str(val)

            if section.chart:
                self._add_chart(doc, section.chart)

            for img in section.images:
                self._add_image(doc, img)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # -- Helper methods --

    @staticmethod
    def _apply_text_style(paragraph: Any, style: Any) -> None:
        """Apply a TextStyle to all runs in a paragraph."""
        if style is None or paragraph is None:
            return
        for run in paragraph.runs:
            if style.font_name:
                run.font.name = style.font_name
            if style.font_size > 0:
                run.font.size = Pt(style.font_size)
            if style.bold:
                run.font.bold = True
            if style.italic:
                run.font.italic = True
            if style.color:
                hex_color = style.color.lstrip("#")
                if len(hex_color) == 6:
                    run.font.color.rgb = RGBColor(
                        int(hex_color[0:2], 16),
                        int(hex_color[2:4], 16),
                        int(hex_color[4:6], 16),
                    )

    @staticmethod
    def _add_callout(doc: Any, text: str) -> None:
        """Add a visually distinct callout paragraph (bold, indented)."""
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        run = para.add_run(f"\u25b6 {text}")
        run.font.bold = True
        run.font.italic = True

    @staticmethod
    def _add_chart(doc: Any, chart_spec: Any) -> None:
        """Render a chart as PNG and embed it in the document."""
        from firefly_dworkers.design.charts import ChartRenderer
        from firefly_dworkers.design.models import DataSeries, ResolvedChart

        # Accept either a ResolvedChart or a dict-like spec
        if isinstance(chart_spec, ResolvedChart):
            resolved = chart_spec
        else:
            # Build a ResolvedChart from the spec
            series = chart_spec.series if hasattr(chart_spec, "series") else []
            if series and isinstance(series[0], dict):
                series = [DataSeries(name=s["name"], values=s["values"]) for s in series]
            resolved = ResolvedChart(
                chart_type=chart_spec.chart_type if hasattr(chart_spec, "chart_type") else chart_spec.get("chart_type", "bar"),
                title=chart_spec.title if hasattr(chart_spec, "title") else chart_spec.get("title", ""),
                categories=chart_spec.categories if hasattr(chart_spec, "categories") else chart_spec.get("categories", []),
                series=series,
                colors=chart_spec.colors if hasattr(chart_spec, "colors") else chart_spec.get("colors", []),
                show_legend=chart_spec.show_legend if hasattr(chart_spec, "show_legend") else chart_spec.get("show_legend", True),
                show_data_labels=chart_spec.show_data_labels if hasattr(chart_spec, "show_data_labels") else chart_spec.get("show_data_labels", False),
                stacked=chart_spec.stacked if hasattr(chart_spec, "stacked") else chart_spec.get("stacked", False),
            )

        renderer = ChartRenderer()
        try:
            png_bytes = renderer._render_to_image_sync(resolved, 800, 600)
            doc.add_picture(io.BytesIO(png_bytes), width=Inches(6))
        except Exception:
            logger.warning("Failed to render chart to image", exc_info=True)

    @staticmethod
    def _add_image(doc: Any, img: Any) -> None:
        """Embed an image from an ImagePlacement spec."""
        if not img.file_path:
            return
        width = Inches(img.width) if img.width else None
        height = Inches(img.height) if img.height else None
        doc.add_picture(img.file_path, width=width, height=height)

    def _modify_sync(self, source: str, operations: list[DocumentOperation]) -> bytes:
        doc = docx.Document(source)

        for op in operations:
            if op.operation == "add_section":
                heading = op.data.get("heading", "")
                content = op.data.get("content", "")
                level = op.data.get("heading_level", 1)
                if heading:
                    doc.add_heading(heading, level=level)
                if content:
                    doc.add_paragraph(content)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
