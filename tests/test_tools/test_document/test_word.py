"""Tests for WordTool adapter."""

from __future__ import annotations

import io
import os
import struct
import tempfile
import zlib

import pytest
from fireflyframework_genai.tools.base import BaseTool

from firefly_dworkers.design.models import ImagePlacement, ResolvedChart, TextStyle
from firefly_dworkers.tools.document.base import DocumentTool
from firefly_dworkers.tools.document.models import SectionSpec
from firefly_dworkers.tools.document.word import WordTool
from firefly_dworkers.tools.registry import tool_registry


def _make_minimal_png(width: int = 1, height: int = 1) -> bytes:
    """Create a minimal valid 1x1 white PNG image in memory."""

    def _chunk(chunk_type: bytes, data: bytes) -> bytes:
        c = chunk_type + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

    header = b"\x89PNG\r\n\x1a\n"
    ihdr_data = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    ihdr = _chunk(b"IHDR", ihdr_data)
    # Single white pixel row: filter byte (0) + RGB (255,255,255)
    raw_row = b"\x00" + b"\xff\xff\xff" * width
    raw_data = raw_row * height
    idat = _chunk(b"IDAT", zlib.compress(raw_data))
    iend = _chunk(b"IEND", b"")
    return header + ihdr + idat + iend


class TestWordToolRegistration:
    def test_is_document_tool(self) -> None:
        assert issubclass(WordTool, DocumentTool)

    def test_is_base_tool(self) -> None:
        assert issubclass(WordTool, BaseTool)

    def test_registry_entry(self) -> None:
        assert tool_registry.has("word")
        assert tool_registry.get_class("word") is WordTool

    def test_category(self) -> None:
        assert tool_registry.get_category("word") == "document"

    def test_name(self) -> None:
        assert WordTool().name == "word"


class TestWordToolRead:
    async def test_read_document(self) -> None:
        docx = pytest.importorskip("docx")

        # Create a minimal .docx in memory
        doc = docx.Document()
        doc.add_heading("Test Title", level=0)
        doc.add_paragraph("Hello world")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(buf.read())
            tmp_path = f.name

        try:
            tool = WordTool()
            result = await tool.execute(action="read", source=tmp_path)
            assert "paragraphs" in result
            assert len(result["paragraphs"]) >= 2
            # First paragraph should be the title heading
            assert result["title"] == "Test Title"
        finally:
            os.unlink(tmp_path)

    async def test_read_detects_headings(self) -> None:
        docx = pytest.importorskip("docx")

        doc = docx.Document()
        doc.add_heading("Chapter 1", level=1)
        doc.add_paragraph("Some content")
        doc.add_heading("Section 1.1", level=2)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(buf.read())
            tmp_path = f.name

        try:
            tool = WordTool()
            result = await tool.execute(action="read", source=tmp_path)
            headings = [p for p in result["paragraphs"] if p["is_heading"]]
            assert len(headings) == 2
            assert headings[0]["heading_level"] == 1
            assert headings[1]["heading_level"] == 2
        finally:
            os.unlink(tmp_path)


class TestWordToolCreate:
    async def test_create_document_basic(self) -> None:
        pytest.importorskip("docx")
        tool = WordTool()
        sections = [
            SectionSpec(heading="Introduction", content="Hello world").model_dump(),
            SectionSpec(heading="Details", bullet_points=["A", "B", "C"]).model_dump(),
        ]
        result = await tool.execute(action="create", title="My Document", sections=sections)
        assert result["success"] is True
        assert result["bytes_length"] > 0

    async def test_create_with_table(self) -> None:
        pytest.importorskip("docx")
        tool = WordTool()
        sections = [
            SectionSpec(
                heading="Data Table",
                table={"headers": ["Name", "Value"], "rows": [["A", "1"]]},
            ).model_dump(),
        ]
        result = await tool.execute(action="create", title="Table Doc", sections=sections)
        assert result["success"] is True

    async def test_create_with_page_break(self) -> None:
        pytest.importorskip("docx")
        tool = WordTool()
        sections = [
            SectionSpec(heading="Page 1", content="First page content").model_dump(),
            SectionSpec(heading="Page 2", content="Second page", page_break_before=True).model_dump(),
        ]
        result = await tool.execute(action="create", title="Multi-page", sections=sections)
        assert result["success"] is True
        assert result["bytes_length"] > 0

    async def test_create_roundtrip(self) -> None:
        """Create a document and then read it back to verify content."""
        pytest.importorskip("docx")

        tool = WordTool()
        # Create via the tool's sync method directly to get bytes
        sections = [SectionSpec(heading="Chapter 1", heading_level=1, content="Body text")]
        data = await tool._create_document("Round Trip", sections)

        # Write to temp file and read back
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(data)
            tmp_path = f.name

        try:
            result = await tool.execute(action="read", source=tmp_path)
            assert result["title"] == "Round Trip"
            heading_texts = [p["text"] for p in result["paragraphs"] if p["is_heading"]]
            assert "Chapter 1" in heading_texts
        finally:
            os.unlink(tmp_path)


class TestWordToolModify:
    async def test_modify_add_section(self) -> None:
        docx = pytest.importorskip("docx")

        # Create a minimal doc
        doc = docx.Document()
        doc.add_heading("Original", level=1)
        doc.add_paragraph("Original content")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(buf.read())
            tmp_path = f.name

        try:
            tool = WordTool()
            result = await tool.execute(
                action="modify",
                source=tmp_path,
                operations=[
                    {
                        "operation": "add_section",
                        "data": {"heading": "New Section", "content": "Added content", "heading_level": 2},
                    }
                ],
            )
            assert result["success"] is True
            assert result["bytes_length"] > 0
        finally:
            os.unlink(tmp_path)


class TestWordToolStyling:
    """Tests for TextStyle application to headings and body text."""

    async def test_heading_style_font_name(self) -> None:
        docx_mod = pytest.importorskip("docx")
        tool = WordTool()
        sections = [
            SectionSpec(
                heading="Styled Heading",
                content="Body text",
                heading_style=TextStyle(font_name="Arial", font_size=24, bold=True),
            ),
        ]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        # The heading is the first paragraph
        heading_para = doc.paragraphs[0]
        assert heading_para.text == "Styled Heading"
        run = heading_para.runs[0]
        assert run.font.name == "Arial"

    async def test_heading_style_bold(self) -> None:
        docx_mod = pytest.importorskip("docx")
        tool = WordTool()
        sections = [
            SectionSpec(
                heading="Bold Heading",
                heading_style=TextStyle(bold=True),
            ),
        ]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        run = doc.paragraphs[0].runs[0]
        assert run.font.bold is True

    async def test_body_style_font(self) -> None:
        docx_mod = pytest.importorskip("docx")
        tool = WordTool()
        sections = [
            SectionSpec(
                content="Styled body text",
                body_style=TextStyle(font_name="Calibri", font_size=14, italic=True),
            ),
        ]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        # The body paragraph is the first paragraph (no heading)
        body_para = doc.paragraphs[0]
        assert body_para.text == "Styled body text"
        run = body_para.runs[0]
        assert run.font.name == "Calibri"
        assert run.font.italic is True

    async def test_body_style_font_size(self) -> None:
        docx_mod = pytest.importorskip("docx")
        from docx.shared import Pt

        tool = WordTool()
        sections = [
            SectionSpec(
                content="Sized text",
                body_style=TextStyle(font_size=16),
            ),
        ]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        run = doc.paragraphs[0].runs[0]
        assert run.font.size == Pt(16)

    async def test_style_color_applied(self) -> None:
        docx_mod = pytest.importorskip("docx")
        from docx.shared import RGBColor

        tool = WordTool()
        sections = [
            SectionSpec(
                content="Colored text",
                body_style=TextStyle(color="#1a73e8"),
            ),
        ]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        run = doc.paragraphs[0].runs[0]
        assert run.font.color.rgb == RGBColor(0x1A, 0x73, 0xE8)

    async def test_heading_and_body_styled_together(self) -> None:
        docx_mod = pytest.importorskip("docx")
        tool = WordTool()
        sections = [
            SectionSpec(
                heading="Title",
                content="Body",
                heading_style=TextStyle(font_name="Georgia", bold=True),
                body_style=TextStyle(font_name="Verdana", italic=True),
            ),
        ]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        heading_run = doc.paragraphs[0].runs[0]
        body_run = doc.paragraphs[1].runs[0]
        assert heading_run.font.name == "Georgia"
        assert body_run.font.name == "Verdana"
        assert body_run.font.italic is True

    async def test_no_style_is_noop(self) -> None:
        """When no style is provided, the tool should still succeed."""
        pytest.importorskip("docx")
        tool = WordTool()
        sections = [
            SectionSpec(heading="Plain", content="No style").model_dump(),
        ]
        result = await tool.execute(action="create", sections=sections)
        assert result["success"] is True

    async def test_apply_text_style_with_none(self) -> None:
        """_apply_text_style should be a no-op when style or paragraph is None."""
        pytest.importorskip("docx")
        # Should not raise
        WordTool._apply_text_style(None, None)
        WordTool._apply_text_style(None, TextStyle())


class TestWordToolNumberedList:
    """Tests for numbered list support."""

    async def test_numbered_list_basic(self) -> None:
        docx_mod = pytest.importorskip("docx")
        tool = WordTool()
        sections = [
            SectionSpec(
                heading="Steps",
                numbered_list=["Step one", "Step two", "Step three"],
            ),
        ]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        # Heading is paragraphs[0]; numbered list items follow
        numbered = [p for p in doc.paragraphs if p.style.name == "List Number"]
        assert len(numbered) == 3
        assert numbered[0].text == "Step one"
        assert numbered[1].text == "Step two"
        assert numbered[2].text == "Step three"

    async def test_numbered_list_empty(self) -> None:
        """Empty numbered list should not add any list paragraphs."""
        docx_mod = pytest.importorskip("docx")
        tool = WordTool()
        sections = [
            SectionSpec(heading="No List", numbered_list=[]),
        ]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        numbered = [p for p in doc.paragraphs if p.style.name == "List Number"]
        assert len(numbered) == 0

    async def test_numbered_and_bullet_lists_coexist(self) -> None:
        docx_mod = pytest.importorskip("docx")
        tool = WordTool()
        sections = [
            SectionSpec(
                bullet_points=["Bullet A", "Bullet B"],
                numbered_list=["Number 1", "Number 2"],
            ),
        ]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
        numbered = [p for p in doc.paragraphs if p.style.name == "List Number"]
        assert len(bullets) == 2
        assert len(numbered) == 2


class TestWordToolCallout:
    """Tests for callout paragraph support."""

    async def test_callout_basic(self) -> None:
        docx_mod = pytest.importorskip("docx")
        from docx.shared import Inches

        tool = WordTool()
        sections = [
            SectionSpec(callout="Important note here"),
        ]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        # Find the callout paragraph (bold + italic with prefix)
        callout_paras = [
            p for p in doc.paragraphs
            if p.runs and p.runs[0].font.bold and p.runs[0].font.italic
        ]
        assert len(callout_paras) == 1
        assert "Important note here" in callout_paras[0].text
        # Check indentation
        assert callout_paras[0].paragraph_format.left_indent == Inches(0.5)

    async def test_callout_empty_is_noop(self) -> None:
        """Empty callout should not add any callout paragraph."""
        docx_mod = pytest.importorskip("docx")
        tool = WordTool()
        sections = [
            SectionSpec(heading="No Callout", callout=""),
        ]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        # Only the heading should be present
        assert len(doc.paragraphs) == 1

    async def test_callout_has_prefix(self) -> None:
        docx_mod = pytest.importorskip("docx")
        tool = WordTool()
        sections = [
            SectionSpec(callout="Check this"),
        ]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        callout_para = doc.paragraphs[0]
        # Should have the triangle prefix
        assert callout_para.text.startswith("\u25b6")


class TestWordToolImages:
    """Tests for image embedding support."""

    async def test_image_embedded(self) -> None:
        docx_mod = pytest.importorskip("docx")
        png_data = _make_minimal_png()
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
            f.write(png_data)
            tmp_path = f.name

        try:
            tool = WordTool()
            sections = [
                SectionSpec(
                    heading="Image Section",
                    images=[ImagePlacement(file_path=tmp_path, width=3.0, height=2.0)],
                ),
            ]
            data = await tool.create(title="", sections=sections)
            # Verify it's a valid docx
            doc = docx_mod.Document(io.BytesIO(data))
            # The document should have inline shapes (images)
            # python-docx adds pictures as inline shapes
            assert len(doc.inline_shapes) >= 1
        finally:
            os.unlink(tmp_path)

    async def test_image_with_no_explicit_size(self) -> None:
        docx_mod = pytest.importorskip("docx")
        png_data = _make_minimal_png()
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
            f.write(png_data)
            tmp_path = f.name

        try:
            tool = WordTool()
            sections = [
                SectionSpec(
                    images=[ImagePlacement(file_path=tmp_path)],
                ),
            ]
            data = await tool.create(title="", sections=sections)
            doc = docx_mod.Document(io.BytesIO(data))
            assert len(doc.inline_shapes) >= 1
        finally:
            os.unlink(tmp_path)

    async def test_multiple_images(self) -> None:
        docx_mod = pytest.importorskip("docx")
        png_data = _make_minimal_png()
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
            f.write(png_data)
            tmp_path = f.name

        try:
            tool = WordTool()
            sections = [
                SectionSpec(
                    images=[
                        ImagePlacement(file_path=tmp_path, width=2.0),
                        ImagePlacement(file_path=tmp_path, width=3.0),
                    ],
                ),
            ]
            data = await tool.create(title="", sections=sections)
            doc = docx_mod.Document(io.BytesIO(data))
            assert len(doc.inline_shapes) >= 2  # noqa: PLR2004
        finally:
            os.unlink(tmp_path)

    async def test_empty_images_list_is_noop(self) -> None:
        docx_mod = pytest.importorskip("docx")
        tool = WordTool()
        sections = [
            SectionSpec(heading="No Images", images=[]),
        ]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        assert len(doc.inline_shapes) == 0

    async def test_image_with_no_file_path_is_skipped(self) -> None:
        docx_mod = pytest.importorskip("docx")
        tool = WordTool()
        sections = [
            SectionSpec(
                images=[ImagePlacement(file_path="", width=2.0)],
            ),
        ]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        assert len(doc.inline_shapes) == 0


class TestWordToolCharts:
    """Tests for chart embedding as PNG images."""

    async def test_chart_embedded_as_image(self) -> None:
        docx_mod = pytest.importorskip("docx")
        pytest.importorskip("matplotlib")
        from firefly_dworkers.design.models import DataSeries

        tool = WordTool()
        chart = ResolvedChart(
            chart_type="bar",
            title="Revenue by Quarter",
            categories=["Q1", "Q2", "Q3"],
            series=[DataSeries(name="2025", values=[100, 200, 300])],
        )
        sections = [
            SectionSpec(heading="Chart Section", chart=chart),
        ]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        # Chart is embedded as an inline picture
        assert len(doc.inline_shapes) >= 1

    async def test_chart_line_type(self) -> None:
        docx_mod = pytest.importorskip("docx")
        pytest.importorskip("matplotlib")
        from firefly_dworkers.design.models import DataSeries

        tool = WordTool()
        chart = ResolvedChart(
            chart_type="line",
            title="Growth Trend",
            categories=["Jan", "Feb", "Mar"],
            series=[DataSeries(name="Revenue", values=[10, 20, 30])],
        )
        sections = [SectionSpec(chart=chart)]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        assert len(doc.inline_shapes) >= 1

    async def test_chart_pie_type(self) -> None:
        docx_mod = pytest.importorskip("docx")
        pytest.importorskip("matplotlib")
        from firefly_dworkers.design.models import DataSeries

        tool = WordTool()
        chart = ResolvedChart(
            chart_type="pie",
            title="Market Share",
            categories=["A", "B", "C"],
            series=[DataSeries(name="Share", values=[45, 30, 25])],
        )
        sections = [SectionSpec(chart=chart)]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        assert len(doc.inline_shapes) >= 1

    async def test_no_chart_when_none(self) -> None:
        docx_mod = pytest.importorskip("docx")
        tool = WordTool()
        sections = [SectionSpec(heading="No Chart", content="Just text")]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        assert len(doc.inline_shapes) == 0

    async def test_chart_multi_series(self) -> None:
        docx_mod = pytest.importorskip("docx")
        pytest.importorskip("matplotlib")
        from firefly_dworkers.design.models import DataSeries

        tool = WordTool()
        chart = ResolvedChart(
            chart_type="bar",
            title="Comparison",
            categories=["Q1", "Q2"],
            series=[
                DataSeries(name="2024", values=[100, 150]),
                DataSeries(name="2025", values=[120, 180]),
            ],
        )
        sections = [SectionSpec(chart=chart)]
        data = await tool.create(title="", sections=sections)
        doc = docx_mod.Document(io.BytesIO(data))
        assert len(doc.inline_shapes) >= 1


class TestWordToolCombined:
    """Tests for combining multiple new features together."""

    async def test_all_features_together(self) -> None:
        docx_mod = pytest.importorskip("docx")
        pytest.importorskip("matplotlib")
        from firefly_dworkers.design.models import DataSeries

        png_data = _make_minimal_png()
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
            f.write(png_data)
            tmp_path = f.name

        try:
            tool = WordTool()
            chart = ResolvedChart(
                chart_type="bar",
                title="Sales",
                categories=["A", "B"],
                series=[DataSeries(name="S1", values=[10, 20])],
            )
            sections = [
                SectionSpec(
                    heading="Full Section",
                    content="Body text here",
                    heading_style=TextStyle(font_name="Georgia", bold=True),
                    body_style=TextStyle(font_name="Verdana", font_size=12),
                    bullet_points=["Bullet A"],
                    numbered_list=["Step 1", "Step 2"],
                    callout="Important note",
                    chart=chart,
                    images=[ImagePlacement(file_path=tmp_path, width=2.0)],
                    table={"headers": ["Col1", "Col2"], "rows": [["a", "b"]]},
                ),
            ]
            data = await tool.create(title="Combined Doc", sections=sections)
            doc = docx_mod.Document(io.BytesIO(data))

            # Title heading
            assert doc.paragraphs[0].text == "Combined Doc"

            # Section heading styled
            section_heading = doc.paragraphs[1]
            assert section_heading.text == "Full Section"
            assert section_heading.runs[0].font.name == "Georgia"

            # Body text styled
            body_para = doc.paragraphs[2]
            assert body_para.text == "Body text here"
            assert body_para.runs[0].font.name == "Verdana"

            # Bullet points
            bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
            assert len(bullets) == 1

            # Numbered list
            numbered = [p for p in doc.paragraphs if p.style.name == "List Number"]
            assert len(numbered) == 2

            # Callout
            callout_paras = [
                p for p in doc.paragraphs
                if p.runs and p.runs[0].font.bold and p.runs[0].font.italic
                and "\u25b6" in p.text
            ]
            assert len(callout_paras) == 1

            # Chart + image = at least 2 inline shapes
            assert len(doc.inline_shapes) >= 2  # noqa: PLR2004
        finally:
            os.unlink(tmp_path)

    async def test_backwards_compatible_basic_create(self) -> None:
        """Existing basic create functionality still works unchanged."""
        pytest.importorskip("docx")
        tool = WordTool()
        sections = [
            SectionSpec(heading="Slide 1", content="Hello world").model_dump(),
            SectionSpec(heading="Slide 2", bullet_points=["A", "B", "C"]).model_dump(),
            SectionSpec(
                heading="Data",
                table={"headers": ["Name", "Value"], "rows": [["X", "1"]]},
            ).model_dump(),
        ]
        result = await tool.execute(action="create", title="Compat Test", sections=sections)
        assert result["success"] is True
        assert result["bytes_length"] > 0

    async def test_styling_with_chart_and_image(self) -> None:
        docx_mod = pytest.importorskip("docx")
        pytest.importorskip("matplotlib")
        from firefly_dworkers.design.models import DataSeries

        png_data = _make_minimal_png()
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
            f.write(png_data)
            tmp_path = f.name

        try:
            tool = WordTool()
            sections = [
                SectionSpec(
                    heading="Styled + Chart",
                    content="Body text",
                    heading_style=TextStyle(font_name="Arial", color="#ff0000"),
                    chart=ResolvedChart(
                        chart_type="line",
                        categories=["X", "Y"],
                        series=[DataSeries(name="S", values=[5, 10])],
                    ),
                    images=[ImagePlacement(file_path=tmp_path)],
                ),
            ]
            data = await tool.create(title="", sections=sections)
            doc = docx_mod.Document(io.BytesIO(data))
            # Heading styled
            heading = doc.paragraphs[0]
            assert heading.runs[0].font.name == "Arial"
            # Chart + image
            assert len(doc.inline_shapes) >= 2  # noqa: PLR2004
        finally:
            os.unlink(tmp_path)


class TestWordToolPublicAPI:
    async def test_create_returns_bytes(self) -> None:
        pytest.importorskip("docx")
        tool = WordTool()
        result = await tool.create(sections=[SectionSpec(heading="H", content="C")])
        assert isinstance(result, bytes)
        assert len(result) > 0

    async def test_create_and_save(self, tmp_path) -> None:
        pytest.importorskip("docx")
        tool = WordTool()
        out = str(tmp_path / "test.docx")
        path = await tool.create_and_save(out, sections=[SectionSpec(heading="Test")])
        assert os.path.exists(path)
        assert os.path.getsize(path) > 0

    async def test_artifact_bytes_after_execute(self) -> None:
        pytest.importorskip("docx")
        tool = WordTool()
        await tool.execute(action="create", sections=[SectionSpec(heading="T").model_dump()])
        assert tool.artifact_bytes is not None
        assert len(tool.artifact_bytes) > 0
